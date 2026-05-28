import logging
import os
from datetime import datetime

from orchestrator.agents.utils.views import print_agent_output
from orchestrator.state import ResearchState

logger = logging.getLogger(__name__)


class PublisherAgent:
    """Assembles the final report string and exports to requested formats.

    Last node in the main graph.  No LLM calls — pure assembly and file I/O.
    """

    def __init__(self, output_dir: str = "./outputs") -> None:
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    async def run(self, state: ResearchState) -> dict:
        task = state["task"]
        title = state.get("title", task["query"])
        date = state.get("date", "")
        introduction = state.get("introduction", "")
        research_data = state.get("research_data", [])
        conclusion = state.get("conclusion", "")
        sources = state.get("sources", [])
        publish_formats = task.get("publish_formats", {"markdown": True})

        print_agent_output("Publishing report...", "PUBLISHER")

        # ── Assemble markdown ─────────────────────────────────────────
        report_parts = [
            f"# {title}",
            f"*{date}*",
            "",
            "## Introduction",
            introduction,
            "",
        ]

        # OPUS FIX (3F): prefer the writer-supplied table_of_contents when present so
        # the writer is authoritative for headers/TOC; fall back to inline assembly
        # for backwards-compat with callers that bypass the writer.
        toc = state.get("table_of_contents") or ""
        if toc:
            report_parts += [toc, ""]
        else:
            report_parts += ["## Table of Contents", ""]
            for i, item in enumerate(research_data, 1):
                anchor = item["section"].lower().replace(" ", "-")
                report_parts.append(f"{i}. [{item['section']}](#{anchor})")
            report_parts.append("")

        # Main sections
        for item in research_data:
            report_parts.append(f"## {item['section']}")
            report_parts.append(item.get("draft", ""))
            report_parts.append("")

        # Conclusion + references
        report_parts += ["## Conclusion", conclusion, "", "## References"]
        if sources:
            report_parts += [f"- {src}" for src in sources]
        else:
            report_parts.append(
                "*Sources are cited inline throughout the report.*"
            )

        full_report = "\n".join(report_parts)

        # ── Save files ────────────────────────────────────────────────
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = (
            "".join(c for c in title if c.isalnum() or c in " -_")[:50].strip()
        )
        base_filename = f"{safe_title}_{timestamp}"
        saved_files: list[str] = []

        if publish_formats.get("markdown", True):
            md_path = os.path.join(self.output_dir, f"{base_filename}.md")
            with open(md_path, "w", encoding="utf-8") as fh:
                fh.write(full_report)
            saved_files.append(md_path)
            print(f"  📄 Markdown: {md_path}")

        if publish_formats.get("docx", False):
            try:
                from docx import Document

                doc = Document()
                doc.add_heading(title, 0)
                for line in full_report.split("\n"):
                    if line.startswith("## "):
                        doc.add_heading(line[3:], level=1)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=2)
                    elif line.strip():
                        doc.add_paragraph(line)
                docx_path = os.path.join(self.output_dir, f"{base_filename}.docx")
                doc.save(docx_path)
                saved_files.append(docx_path)
                print(f"  📝 DOCX: {docx_path}")
            except Exception as exc:
                logger.warning("DOCX export failed: %s", exc)
                print(f"  ⚠️ DOCX export failed: {exc}")

        print_agent_output(
            f"Report published successfully.\n"
            f"Files saved: {len(saved_files)}\n"
            f"Total length: {len(full_report)} chars",
            "PUBLISHER",
        )

        return {"report": full_report}
